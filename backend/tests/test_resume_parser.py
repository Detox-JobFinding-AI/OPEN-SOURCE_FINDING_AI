import io
import os
import pytest
import docx
import pdfplumber
from app.utils.resume_parser import (
    extract_text,
    extract_text_from_pdf,
    extract_text_from_docx,
    extract_text_from_txt,
)

# Helper to generate a valid in-memory DOCX file with optional header and footer
def create_mock_docx(paragraphs, table_contents=None, header_text=None, footer_text=None, duplicate_sections=False) -> bytes:
    doc = docx.Document()
    
    if duplicate_sections:
        doc.add_section()
        
    for section in doc.sections:
        if header_text:
            header = section.header
            header.paragraphs[0].text = header_text
        if footer_text:
            footer = section.footer
            footer.paragraphs[0].text = footer_text
            
    for p in paragraphs:
        doc.add_paragraph(p)
        
    if table_contents:
        table = doc.add_table(rows=len(table_contents), cols=len(table_contents[0]))
        for r_idx, row in enumerate(table_contents):
            for c_idx, val in enumerate(row):
                table.cell(r_idx, c_idx).text = val
                
    stream = io.BytesIO()
    doc.save(stream)
    return stream.getvalue()

# Mock structures for pdfplumber
class MockPage:
    def __init__(self, text, raise_err=False):
        self._text = text
        self._raise_err = raise_err
    def extract_text(self):
        if self._raise_err:
            raise Exception("Corrupt PDF vector graphic on page")
        return self._text

class MockPDF:
    def __init__(self, pages):
        self.pages = pages
    def __enter__(self):
        return self
    def __exit__(self, exc_type, exc_val, exc_tb):
        pass

# --- TXT Extraction Tests ---

def test_extract_text_from_txt_success():
    content = "Name: John Doe\nRole: Python Developer\nLocation: New York\nAbout: Experience in writing modular utilities."
    file_bytes = content.encode("utf-8")
    extracted = extract_text_from_txt(file_bytes)
    assert extracted == content

def test_extract_text_from_txt_latin1_fallback():
    content = "Café Candidate resume details and experience in backend systems"
    file_bytes = content.encode("latin-1")
    extracted = extract_text_from_txt(file_bytes)
    assert "Café" in extracted

def test_extract_text_from_txt_utf8_bom_stripped():
    # Test that UTF-8 BOM is stripped automatically by utf-8-sig
    content = "Name: John Doe"
    bom_content = b"\xef\xbb\xbf" + content.encode("utf-8")
    extracted = extract_text_from_txt(bom_content)
    
    assert extracted == content
    assert not extracted.startswith("\ufeff")  # BOM should be stripped

def test_extract_text_from_txt_empty():
    extracted = extract_text_from_txt(b"")
    assert extracted == ""

# --- DOCX Extraction Tests ---

def test_extract_text_from_docx_success():
    paragraphs = ["Jane Doe", "Full Stack Developer", "Experience in modern web frameworks"]
    tables = [["Skill", "Rating"], ["React", "Expert"]]
    docx_bytes = create_mock_docx(
        paragraphs=paragraphs,
        table_contents=tables,
        header_text="Contact: jane@example.com",
        footer_text="Page 1 of 2"
    )
    
    extracted = extract_text_from_docx(docx_bytes)
    assert "Jane Doe" in extracted
    assert "Full Stack Developer" in extracted
    assert "Skill" in extracted
    assert "React" in extracted
    assert "Contact: jane@example.com" in extracted
    assert "Page 1 of 2" in extracted

def test_extract_text_from_docx_duplicate_headers_prevented():
    paragraphs = ["Jane Doe Resume Details"]
    docx_bytes = create_mock_docx(
        paragraphs=paragraphs,
        header_text="DUPLICATE HEADER TEXT",
        footer_text="DUPLICATE FOOTER TEXT",
        duplicate_sections=True
    )
    
    extracted = extract_text_from_docx(docx_bytes)
    assert extracted.count("DUPLICATE HEADER TEXT") == 1
    assert extracted.count("DUPLICATE FOOTER TEXT") == 1

def test_extract_text_from_docx_corrupted():
    with pytest.raises(ValueError, match="Failed to extract text from DOCX"):
        extract_text_from_docx(b"invalid corrupt zip content")

# --- PDF Extraction Tests ---

def test_extract_text_from_pdf_success(monkeypatch):
    pages = [
        MockPage("John Doe Resume Page 1: This is some dummy text to exceed the fifty character threshold required by the parser."),
        MockPage("Experience details Page 2: This is some dummy text to exceed the fifty character threshold required by the parser.")
    ]
    monkeypatch.setattr(pdfplumber, "open", lambda f: MockPDF(pages))
    
    extracted = extract_text_from_pdf(b"mock pdf bytes")
    assert "John Doe Resume Page 1" in extracted
    assert "Experience details Page 2" in extracted

def test_extract_text_from_pdf_page_fault_tolerance(monkeypatch):
    pages = [
        MockPage("Corrupt page text", raise_err=True),
        MockPage("Page 2 Text: This is some dummy text to exceed the fifty character threshold required by the parser.")
    ]
    monkeypatch.setattr(pdfplumber, "open", lambda f: MockPDF(pages))
    
    extracted = extract_text_from_pdf(b"mock pdf bytes")
    assert "Page 2 Text" in extracted
    assert "Corrupt page text" not in extracted

def test_extract_text_from_pdf_scanned_fails(monkeypatch):
    pages = [MockPage(""), MockPage("   ")]
    monkeypatch.setattr(pdfplumber, "open", lambda f: MockPDF(pages))
    with pytest.raises(ValueError, match="This PDF may be scanned or image-only"):
        extract_text_from_pdf(b"mock scanned pdf")

def test_extract_text_from_pdf_low_character_count_fails(monkeypatch):
    pages = [MockPage("Resume"), MockPage("Page 2")]
    monkeypatch.setattr(pdfplumber, "open", lambda f: MockPDF(pages))
    with pytest.raises(ValueError, match="Extracted text is too short"):
        extract_text_from_pdf(b"mock scanned pdf")

def test_extract_text_from_pdf_empty(monkeypatch):
    monkeypatch.setattr(pdfplumber, "open", lambda f: MockPDF([]))
    with pytest.raises(ValueError, match="The PDF file contains no pages"):
        extract_text_from_pdf(b"mock empty pdf")

def test_extract_text_from_pdf_corrupted(monkeypatch):
    def mock_raise(*args, **kwargs):
        raise Exception("Corrupt PDF structure")
    monkeypatch.setattr(pdfplumber, "open", mock_raise)
    
    with pytest.raises(ValueError, match="Failed to extract text from PDF"):
        extract_text_from_pdf(b"corrupt pdf content")

# --- Stream Type and Rejection Tests ---

def test_extract_text_rejects_string_io():
    string_io = io.StringIO("some text data")
    with pytest.raises(TypeError, match="Text-mode streams are not supported"):
        extract_text(string_io, "resume.txt")

def test_extract_text_rejects_text_file(tmp_path):
    # Create a text file opened in text mode 'r'
    temp_file = tmp_path / "temp_resume.txt"
    temp_file.write_text("file text data", encoding="utf-8")
    
    with open(temp_file, "r", encoding="utf-8") as f:
        with pytest.raises(TypeError, match="Text-mode streams are not supported"):
            extract_text(f, "temp_resume.txt")

# --- Unified extract_text Routing Tests ---

def test_extract_text_routing_pdf(monkeypatch):
    pages = [
        MockPage("PDF Extracted Text: This is some dummy text to exceed the fifty character threshold required by the parser.")
    ]
    monkeypatch.setattr(pdfplumber, "open", lambda f: MockPDF(pages))
    
    res = extract_text(b"pdf data", "resume.pdf")
    assert res == pages[0]._text

def test_extract_text_routing_pdf_with_query_params(monkeypatch):
    pages = [
        MockPage("PDF Extracted Text: This is some dummy text to exceed the fifty character threshold required by the parser.")
    ]
    monkeypatch.setattr(pdfplumber, "open", lambda f: MockPDF(pages))
    
    res = extract_text(b"pdf data", "  resume.pdf?version=12&user=3#anchor ")
    assert res == pages[0]._text

def test_extract_text_routing_docx():
    docx_bytes = create_mock_docx(["DOCX Text: This is some dummy text to exceed the fifty character threshold required by the parser."])
    res = extract_text(docx_bytes, "my_resume.docx")
    assert "DOCX Text" in res

def test_extract_text_routing_txt():
    res = extract_text(b"TXT Text", "resume.txt")
    assert res == "TXT Text"

def test_extract_text_unsupported_doc_extension():
    with pytest.raises(ValueError, match="Unsupported format: Legacy Word '.doc'"):
        extract_text(b"some data", "resume.doc")

def test_extract_text_unsupported_other_extension():
    with pytest.raises(ValueError, match="Unsupported file extension '.png'"):
        extract_text(b"some data", "resume.png")

def test_extract_text_missing_filename():
    with pytest.raises(ValueError, match="Filename must be provided"):
        extract_text(b"some data", "")

def test_extract_text_file_not_found():
    with pytest.raises(FileNotFoundError):
        extract_text("non_existent_file.pdf", "non_existent_file.pdf")

def test_extract_text_directory_fails():
    current_dir = os.path.dirname(os.path.abspath(__file__))
    with pytest.raises(FileNotFoundError, match="File not found or is a directory"):
        extract_text(current_dir, "some_dir.pdf")

def test_extract_text_stream_cursor_reset(monkeypatch):
    pages = [
        MockPage("PDF Extracted Text: This is some dummy text to exceed the fifty character threshold required by the parser.")
    ]
    monkeypatch.setattr(pdfplumber, "open", lambda f: MockPDF(pages))
    
    stream = io.BytesIO(b"pdf data")
    stream.read()  # Move cursor to end of stream
    assert stream.tell() > 0
    
    res = extract_text(stream, "resume.pdf")
    assert res == pages[0]._text
    assert stream.tell() == 0  # Verify cursor was reset back to 0

def test_extract_text_default_cleaning_behavior(monkeypatch):
    pages = [
        MockPage("PDF Extracted   Text:   This is some   dummy text to exceed the fifty character threshold required by the parser.\n•   Some bullet point.")
    ]
    monkeypatch.setattr(pdfplumber, "open", lambda f: MockPDF(pages))
    
    res = extract_text(b"pdf data", "resume.pdf")
    # Spaces collapsed, bullet normalized
    assert "PDF Extracted Text: This is some dummy text" in res
    assert "- Some bullet point." in res

def test_extract_text_clean_false_bypasses_cleaning(monkeypatch):
    pages = [
        MockPage("PDF Extracted   Text:   This is some   dummy text to exceed the fifty character threshold required by the parser.\n•   Some bullet point.")
    ]
    monkeypatch.setattr(pdfplumber, "open", lambda f: MockPDF(pages))
    
    res = extract_text(b"pdf data", "resume.pdf", clean=False)
    # Verify raw text is returned, including double spaces and uncleaned bullet symbol
    assert "PDF Extracted   Text:   This is some   dummy text" in res
    assert "•   Some bullet point." in res

