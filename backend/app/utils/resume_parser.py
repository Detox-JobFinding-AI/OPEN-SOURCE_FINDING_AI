import contextlib
import io
import os
from typing import Generator, Union, BinaryIO
import pdfplumber
import docx

@contextlib.contextmanager
def _open_file_input(file_input: Union[str, bytes, BinaryIO]) -> Generator[BinaryIO, None, None]:
    """
    Context manager to yield a BinaryIO stream from either a file path (str),
    raw bytes, or an existing BinaryIO stream. Handles resource closing if opened.
    
    Rejects text-mode streams.
    """
    # Reject explicit text-mode stream classes
    if isinstance(file_input, (io.TextIOBase, io.StringIO)):
        raise TypeError("Invalid file input: Text-mode streams are not supported. Must be binary.")
        
    if isinstance(file_input, str):
        if not os.path.isfile(file_input):
            raise FileNotFoundError(f"File not found or is a directory: {file_input}")
        with open(file_input, 'rb') as f:
            yield f
    elif isinstance(file_input, bytes):
        yield io.BytesIO(file_input)
    elif hasattr(file_input, 'read'):
        # Reject duck-typed text-mode file objects
        if getattr(file_input, 'encoding', None) is not None:
            raise TypeError("Invalid file input: Text-mode streams are not supported. Must be binary.")
            
        if hasattr(file_input, 'seek'):
            try:
                file_input.seek(0)
            except Exception:
                pass
        yield file_input
    else:
        raise TypeError("Invalid file input type. Must be str, bytes, or a file-like stream.")

def extract_text_from_pdf(file_input: Union[str, bytes, BinaryIO]) -> str:
    """
    Extracts text from a PDF file using pdfplumber.
    Raises ValueError if PDF is corrupted or scanned (no text extracted).
    """
    MIN_PDF_CHARACTERS = 50  # Production threshold to detect scanned or low-text PDFs
    
    try:
        with _open_file_input(file_input) as stream:
            with pdfplumber.open(stream) as pdf:
                if not pdf.pages:
                    raise ValueError("The PDF file contains no pages.")
                text_pages = []
                for page in pdf.pages:
                    # Page-level extraction fault tolerance
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text_pages.append(page_text)
                    except Exception:
                        # Skip failed page, preserving the rest of the document
                        pass
                
                combined_text = "\n".join(text_pages)
                # Check against a minimum character threshold to identify scanned PDFs
                if len(combined_text.strip()) < MIN_PDF_CHARACTERS:
                    raise ValueError(
                        f"Extracted text is too short ({len(combined_text.strip())} characters). "
                        "This PDF may be scanned or image-only and requires OCR."
                    )
                return combined_text
    except Exception as e:
        if isinstance(e, (FileNotFoundError, TypeError, ValueError)):
            raise e
        raise ValueError(f"Failed to extract text from PDF: {str(e)}") from e

def extract_text_from_docx(file_input: Union[str, bytes, BinaryIO]) -> str:
    """
    Extracts text from a DOCX file using python-docx, including headers, footers, body, and tables.
    """
    try:
        with _open_file_input(file_input) as stream:
            doc = docx.Document(stream)
            paragraphs = []
            seen_header_footers = set()
            
            # Extract text from headers and footers with duplicate prevention
            for section in doc.sections:
                if section.header:
                    for p in section.header.paragraphs:
                        p_text = p.text.strip()
                        if p_text and p_text not in seen_header_footers:
                            seen_header_footers.add(p_text)
                            paragraphs.append(p.text)
                if section.footer:
                    for p in section.footer.paragraphs:
                        p_text = p.text.strip()
                        if p_text and p_text not in seen_header_footers:
                            seen_header_footers.add(p_text)
                            paragraphs.append(p.text)
            
            # Extract body paragraphs
            for p in doc.paragraphs:
                if p.text.strip():
                    paragraphs.append(p.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text)
                            
            return "\n".join(paragraphs)
    except Exception as e:
        if isinstance(e, (FileNotFoundError, TypeError)):
            raise e
        raise ValueError(f"Failed to extract text from DOCX: {str(e)}") from e

def extract_text_from_txt(file_input: Union[str, bytes, BinaryIO]) -> str:
    """
    Extracts text from a plain TXT file, trying UTF-8-sig first to handle BOM, falling back to Latin-1.
    """
    try:
        with _open_file_input(file_input) as stream:
            raw_bytes = stream.read()
            try:
                # Use utf-8-sig to automatically strip BOM characters if present
                return raw_bytes.decode('utf-8-sig')
            except UnicodeDecodeError:
                return raw_bytes.decode('latin-1')
    except Exception as e:
        if isinstance(e, (FileNotFoundError, TypeError)):
            raise e
        raise ValueError(f"Failed to extract text from TXT: {str(e)}") from e

def extract_text(file_input: Union[str, bytes, BinaryIO], filename: str) -> str:
    """
    Unified extraction routing function. Identifies format from the filename
    extension and delegates to the appropriate extractor.
    
    Supported formats: .pdf, .docx, .txt
    """
    if not filename:
        raise ValueError("Filename must be provided to determine the file type.")
    
    clean_filename = filename.strip().split('?')[0].split('#')[0]
    
    ext = os.path.splitext(clean_filename.lower())[1]
    if ext == '.pdf':
        return extract_text_from_pdf(file_input)
    elif ext == '.docx':
        return extract_text_from_docx(file_input)
    elif ext == '.txt':
        return extract_text_from_txt(file_input)
    elif ext == '.doc':
        raise ValueError("Unsupported format: Legacy Word '.doc' files are not supported. Please convert to '.docx' or '.pdf'.")
    else:
        raise ValueError(f"Unsupported file extension '{ext}'. Only .pdf, .docx, and .txt files are supported.")
